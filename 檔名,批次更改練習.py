# 此程式碼 功能 將資料夾包含子目錄 有jpg都會複製到output2資料夾並命名p01 , p02 序列檔名
import hashlib
import os
import shutil
import re


def imgname():
    # 指定資料夾 ↓
    path_ls = r'C:/圖片測試'
    sample_tree = os.walk(path_ls)  # 以os.walk() 搜尋目前路徑及其子目錄。

    # 如果要使用 py當前路徑
    # cur_path=os.path.dirname(__file__) # 取得目前路徑
    # sample_tree=os.walk(cur_path) #以os.walk() 搜尋目前路徑及其子目錄。

    output_dir = 'output2'  # 儲存檔案的資料為output2

    # ↓ 讀取資料夾名稱、下一層資料串列和資料夾中所有檔案串列。
    # os.walk 返回的是一個三元組(root,dirs,files)。
    # root 所指的是當前正在遍歷的這個文件夾的本身的地址
    # dirs 是一個 list ，內容是該文件夾中所有的目錄的名字(不包括子目錄)
    # files 同樣是 list , 內容是該文件夾中所有的文件(不包括子目錄)
    for dirname, subdir, files in sample_tree:
        allfiles = []
    # os.path.basename(path) 	返回文件名 os.path.basename('c:\test.csv') 結果 'test.csv' 或 抓取資料夾名稱
    # 例如 c:/anaconda\爬蟲阿摩考古題 會取得到 爬蟲阿摩考古題
    # os.path.basename(path) 	返回文件名 os.path.basename('c:\test.csv') 結果 'test.csv'
        basename = os.path.basename(dirname)

    # 如果是output2 資料夾不必處理，因為 下面os.path.exists(target_dir) 會產生 output2 資料夾,
    # 並將.jpg 複製到output2 資料夾中,
    # 為了避免第二次執行由子output2 進行處理，因此在第9列設定 output2 資料夾不予處理。
        if basename == output_dir:  # 注意執行是在for內， 有output2 目錄不再重複處理
            continue

        for file in files:  # 讀取所有 jpg 檔名，存入 allfiles 串列中
            # split分割 例如檔名：ABC.jpg 從.  切下去 type是list 抓取尾巴的jpg index 是 [-1]
            ext = file.split('.')[-1]
            if ext == "jpg":  # 讀取 *.jpg to allfiles  如果找到jpg 加入到list ，將檔案的副檔名存至ext 變數中。
                allfiles.append(file)

        if len(allfiles) > 0:  # 並確認有檔案才處理，避免產生誤動作。
            target_dir = dirname + '/' + output_dir  # 組合dirname本身遍歷地址路徑 + 資料夾名稱
            if not os.path.exists(target_dir):  # 如果output2 資料夾未建立，建立 output2 資料夾。
                os.mkdir(target_dir)  # 建立 output2 資料夾

        # 將 jpg 存入 output2 目錄中
        # 逐一複製所有的 jpg 檔，來源檔為 srcfile，是以 srcfile=dirname +"/" + file 組成完整路徑名稱。
        # destfile 是目標檔案，是由 target_dir + "/" + m_filename + '.jpg'
        # 組成完整路徑名稱，即 <output2/p0.jpg>、<output2/p1.jpg>... 複製檔案。
            counter = 0
            for file in allfiles:
                # filename=file.split('.')[0] #主檔名 不抓副檔名 沒有用意 是程式原作者 寫爽的
                # print(filename)
                m_filename = "p" + str(counter)  # 取檔案名稱 p01  p02
                destfile = target_dir + '/' + m_filename + '.jpg'  # 加上完整路徑
                srcfile = dirname + "/" + file
                print(destfile)
                shutil.copy(srcfile, destfile)  # 複製檔案
                counter += 1

        print("完成...")


# import os,shutil


def mp3copy():  # 副檔名mp3 去除有不合法字元
    output_dir = 'output'
    path_ls = r'C:/圖片測試'
    sample_tree = os.walk(path_ls)
    # cur_path=os.path.dirname(__file__) # 取得目前路徑

    for dirname, subdir, files in sample_tree:
        allfiles = []
        basename = os.path.basename(dirname)  # 返回檔案名稱
        if basename == output_dir:  # 存在output 目錄不再重複處理
            continue

        for file in files:  # 讀取所有 mp3 檔名，存入 allfiles 串列中

            # file.split('.')[-1]從 . 分割 例如 my.mp3  變成list的資料 ['my', 'mp3'] 只抓取mp3
            ext = file.split('.')[-1]
            if ext == "mp3":  # 讀取 *.mp3 to allfiles  過濾功能 有mp3的副檔名 加入到allfiles的list中
                allfiles.append(file)

        if len(allfiles) > 0:  # 將 mp3 存入 output 目錄中
            target_dir = dirname + '/' + output_dir
            if not os.path.exists(target_dir):
                os.mkdir(target_dir)

            for file in allfiles:
                # 主檔名 檔名分割 例如 my.mp3  變成list的資料 ['my', 'mp3'] 只抓取my
                filename = file.split('.')[0]
                fileExtension = file.split('.')[-1]
                m_filename = ""
                # print(filename)

                # 第一種過濾方法
                # 將主檔名中不合法的字元去除 如果檔名為 (smskf，.jajd ias) 會一個一個字去找 找到並加變為空
                for c in filename:
                    if c == " " or c == "." or c == "," or c == "、" or c == "，" or c == "(" or c == ")":
                        m_filename += ""  # 去除不合法字元
                        # print(m_filename)
                    else:
                        m_filename += c
                # 以下#為第一種過濾方法 需要在解開 使用第二種最快速
                # destfile = "{}.{}".format(
                #     target_dir+'/'+m_filename, fileExtension)  # 加上完整路徑
                # srcfile = dirname + "/" + file
                # # print(srcfile)
                # print(destfile)
                # shutil.copy(srcfile, destfile)  # 複製檔案到output 資料夾

                # 第二種過濾方法
                b_filename = ""
                for i in re.split("-|、|'| |,", filename):
                    b_filename += i
                # print(b_filename)
                destfile = "{}.{}".format(
                    target_dir+'/'+b_filename, fileExtension)  # 加上完整路徑
                srcfile = dirname + "/" + file
                # print(srcfile)
                print(destfile)
                shutil.copy(srcfile, destfile)  # 複製檔案到output 資料夾
    print("完成...")


def FindSamePic():  # 功能找出相同圖片png jpg
    path_ls = r'C:/圖片測試'
    # cur_path = os.path.dirname(__file__)  # 取得目前路徑

    allmd5s = dict()  # 利用字典儲存為第一張發現的圖片，以便可以MD5 呼叫相同的MD5

    n = 0
    # dirname 當前路徑, subdir所有資料夾路徑, files所有的檔案名稱
    for dirname, subdir, files in os.walk(path_ls):
        allfiles = []
        for file in files:  # 取得所有 .png .jpg 檔，存入 allfiles 串列中
            ext = file.split('.')[-1]  # 取得檔案名稱的副檔名
            if ext == "png" or ext == "jpg":  # 過濾 只找 png及jpg
                allfiles.append(dirname + '/'+file)  # 加入到allfiles的list

        # 逐一取出照片檔並和其目錄下(包含所有子目錄)的所有照片，比較是否相同。
        if len(allfiles) > 0:  # 若有資料執行
            for imagefile in allfiles:  # allfiles 找到的文件檔案的名稱路徑
                img_md5 = hashlib.md5(
                    open(imagefile, 'rb').read()).digest()  # 開啟文件取得文件md5
                # print(img_md5)
                # 如果 img_md5 編碼存在 allmd5s 串列中，表示該張照片重複。 否則將 img_md5 加入 allmd5s 串列中。
                # 思考 有兩張相同  先取得第一張md5 第二張md5一定比對的到相同的
                if img_md5 in allmd5s:  # img_md5 資料有沒有 在 allmd5s裡面 有會是True 沒有會false
                    if n == 0:  # 用意 只顯示一次"找到下列重覆的檔案："
                        print("找到下列重覆的檔案：")
                    n += 1
                    print(os.path.abspath(imagefile))
                    print(allmd5s[img_md5] + "\n")  # print 字典中的md5
                    print(allmd5s)
                else:
                    allmd5s[img_md5] = os.path.abspath(
                        imagefile)  # 組合字典 {md5:檔案路徑}
                    # print(allmd5s)
    print("完成...")


def photoReSize():  # 更改圖檔為相同的大小
    import os
    from PIL import Image

    image_dir = 'resized_photo'
    image_width = 800

    path_ls = r'C:/圖片測試'
    # cur_path = os.path.dirname(__file__)  # 取得目前路徑
    # sample_tree = os.walk(path_ls)

    # 返回的資料型態 dirname目前資料夾路徑, subdir所有資料夾路徑, files所有檔案名稱
    for dirname, subdir, files in os.walk(path_ls):
        allfiles = []

        # os.path.basename(path) 	返回文件名 os.path.basename('c:\test.csv') 結果 'test.csv'
        basename = os.path.basename(dirname)
        if basename == image_dir:  # resized_photo 目錄不再重複處理，找當前路徑下的所有資料夾 如果有image_dir = 'resized_photo' 使用continue
            continue
        for file in files:  # 取得所有 .png .jpg 檔，存入 allfiles 串列中
            ext = file.split('.')[-1]  # aa.jpg 使用split('.')[-1]分割 取得副檔名
            if ext == "png" or ext == "jpg":  # 過濾 只抓取有jpg及png的檔案
                allfiles.append(dirname + '/'+file)  # 組合路徑 加入到list中

        if len(allfiles) > 0:  # list有資料才執行
            # 當前路徑+ image_dir = resized_photo 儲存的資料夾
            target_dir = dirname + '/' + image_dir

            # 如果沒有resized_photo 就建立resized_photo
            if not os.path.exists(target_dir):
                os.mkdir(target_dir)
            for file in allfiles:  # 跑 allfiles過濾好的圖片路徑
                # os.path.split 如果給出的是一個目錄和文件名，則輸出路徑和文件名
                # os.path.split('C:/soft/python/test.py') : ('C:/soft/python', 'test.py')
                # os.path.split('C:/soft/python/test') : ('C:/soft/python', 'test')
                # os.path.split('C:/soft/python/') : ('C:/soft/python', '')
                pathname, filename = os.path.split(
                    file)  # 共取得兩個值
                img = Image.open(file)  # 開啟圖片
                w, h = img.size  # 該圖片大小
                print(w, h)
                # 注意 image_width 變數設定 為 800
                # 值1 = 寬度 值2 = 高度 800/w*h 除高算出適當的高度 寬度算出比例高度公式(高/高*寬)
                img = img.resize((image_width, int(image_width/float(w)*h)))
                img.save(target_dir+'/'+filename)  # 儲存圖
                print("<{}> 複製完成!".format(target_dir+'/'+filename))
                img.close()  # 關閉圖

    print("完成...")


def FindKeyWord():
    import os
    path_ls = r'C:/圖片測試'
    cur_path = os.path.dirname(__file__)  # 取得目前路徑
    sample_tree = os.walk(path_ls)  # 使用os.walk 取得三個值 如下面
    keyword = "shutil"

    # 返回的資料型態 dirname目前資料夾路徑, subdir所有資料夾路徑, files所有檔案名稱
    for dirname, subdir, files in sample_tree:
        allfiles = []
        # print(dirname)
        for file in files:  # 取得所有 .py .txt 檔，存入 allfiles 串列中
            ext = file.split('.')[-1]  # 讀取檔案副檔名
            if ext == "py" or ext == "txt":  # 過濾 只找py 及 txt
                allfiles.append(dirname + '/'+file)  # 加入並組合路徑
            # print(allfiles)
                # print(allfiles)

        if len(allfiles) > 0:
            for file in allfiles:  # 讀取 allfiles 串列所有檔案
                # print(file)
                # 逐一讀取所有檔後一列一列比對是否找到指定的字元，讀取檔案最 好以 try.except 補捉錯誤,以防程式中止。
                try:
                    fp = open(file, "r", encoding='UTF-8')
                    article = fp.readlines()  # 將內容儲存在article
                    fp.close
                    line = 0
                    for row in article:
                        line += 1  # 每次搜尋一行 +1
                        if keyword in row:
                            print(f"在 {file}，第 {line} 列找到「{keyword}」。")
                except:
                    print("{} 無法讀取..." .format(file))

    print("完成...")


def FindKeyWord2():
    import os
    import docx

    path_ls = r'C:/圖片測試'
    cur_path = os.path.dirname(__file__)  # 取得目前路徑
    sample_tree = os.walk(path_ls)

    keyword = "籃球"
    print("搜尋字串：{}" .format(keyword))

    for dirname, subdir, files in sample_tree:  # 與以上所有程式碼大同小異
        allfiles = []
        for file in files:  # 取得所有 .docx 檔，存入 allfiles 串列中
            ext = file.split('.')[-1]
            if ext == "docx":  # get *.docx to allfiles
                allfiles.append(dirname + '/'+file)  # 組合路徑

        for file in allfiles:
            print("正在搜尋 <{}> 檔案...".format(file))
            try:
                doc = docx.Document(file)  # 使用docx庫開啟
                line = 0
                for p in doc.paragraphs:  # 注意doc.paragraphs 使用docx的函式 paragraphs意思是段落 讀取段落
                    line += 1
                    if keyword in p.text:  # 一行 一行的讀取 如果p.text 裡面有keyword 則True
                        print("...在第 {} 段文字中找到「{}」\n {}。".format(
                            line, keyword, p.text))
            except:
                print("無法讀取 {} 檔案..." .format(file))

    print("\n搜尋完畢...")


def FindKeyWord3():
    import os
    import docx
    import sys
    # 匯入sys 模組，並以 len(sys.argv)取得參數的個數，如果我們是在 Anaconda Prompt 視窗執行「python FindKeyWord3.py 籃 球」指令，
    # 則參數的個數將是2，索引第0、1個參數分別是 「FindKeyWord3.py」和「籃球」。

    if len(sys.argv) == 1:  # 如果只輸入 python me.py 不加關鍵字 才會是返回1
        keyword = "shutil"
        print("語法：python FindKeyWord3.py 搜尋字串\n")
    else:
        keyword = sys.argv[1]  # 例如 設定索引第1個參數為「籃球」搜尋字串。 這樣len(sys.argv)長度等於2

    # cur_path=os.path.dirname(__file__) # 取得目前路徑
    path_ls = r'C:/圖片測試'
    cur_path = os.getcwd()  # 改用os.getcwd()取得現在的路徑名稱。
    sample_tree = os.walk(path_ls)
    print(cur_path)

    for dirname, subdir, files in sample_tree:  # 值1 當前路徑 值2 所有資料夾名稱 值3 所有文件名稱
        allfiles = []
        for file in files:  # 取得所有 .py .txt .docx 檔，存入 allfiles 串列中
            ext = file.split('.')[-1]
            if ext == "py" or ext == "txt" or ext == "docx":
                allfiles.append(dirname + '/'+file)  # 組合路徑

        if len(allfiles) > 0:
            for file in allfiles:  # 讀取 allfiles 串列所有檔案
                try:
                    if file.split('.')[-1] == "docx":  # .docx
                        doc = docx.Document(file)  # 讀取docx
                        line = 0
                        for p in doc.paragraphs:  # 讀取每行 一行一行的讀
                            line += 1  # 可得知道讀到第幾行
                            if keyword in p.text:
                                print("...在第 {} 段文字中找到「{}」\n {}。".format(
                                    line, keyword, p.text))
                    else:  # .py or .txt #如果非docx就會執行此區
                        fp = open(file, "r", encoding='UTF-8')
                        article = fp.readlines()
                        fp.close
                        line = 0
                        for row in article:
                            line += 1
                            if keyword in row:
                                print("在 {}，第 {} 列找到「{}」。".format(
                                    file, line, keyword))
                except:
                    print("{} 無法讀取..." .format(file))

    print("完成...")


if __name__ == "__main__":
    # imgname()
    # mp3copy()
    # FindSamePic()
    # photoReSize()
    # FindKeyWord()
    # FindKeyWord2()
    # FindKeyWord3()
    